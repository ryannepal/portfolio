# Builds "Portfolio Editing Guide.docx" — a real Word document with proper
# heading styles, tables and shaded code blocks. Run:  py build_guide.py
# This script is a build tool, not part of the website. Safe to delete.

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ACCENT = RGBColor(0x1B, 0x4D, 0xD1)
INK = RGBColor(0x1A, 0x18, 0x14)
MUTE = RGBColor(0x5E, 0x57, 0x49)
CODE_BG = "F2EFE8"
NOTE_BG = "EDF1FB"

doc = Document()

# ---------------------------------------------------------------- base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ("Title", 30, INK, 0, 6),
    ("Heading 1", 19, ACCENT, 22, 8),
    ("Heading 2", 14, INK, 14, 5),
    ("Heading 3", 12, MUTE, 11, 4),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.font.italic = False
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for m in ("top", "bottom", "left", "right"):
    setattr(doc.sections[0], f"{m}_margin", Inches(0.9))


# --------------------------------------------------------------------- helpers
def shade(paragraph, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    paragraph._p.get_or_add_pPr().append(el)


def border(paragraph, hex_color, side="left", size=18):
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "8")
    b.set(qn("w:color"), hex_color)
    pbdr.append(b)
    paragraph._p.get_or_add_pPr().append(pbdr)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text="", bold=False, italic=False, size=None, color=None, space_after=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if space_after is not None:
        par.paragraph_format.space_after = Pt(space_after)
    return par


def rich(parts, style=None):
    """parts: list of (text, kind) where kind is '', 'b', 'i' or 'code'."""
    par = doc.add_paragraph(style=style)
    for text, kind in parts:
        run = par.add_run(text)
        if kind == "b":
            run.bold = True
        elif kind == "i":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
            run.font.color.rgb = ACCENT
    return par


def steps(items):
    for item in items:
        par = doc.add_paragraph(style="List Number")
        if isinstance(item, str):
            par.add_run(item)
        else:
            for text, kind in item:
                run = par.add_run(text)
                if kind == "b":
                    run.bold = True
                elif kind == "code":
                    run.font.name = "Consolas"
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = ACCENT
        par.paragraph_format.space_after = Pt(4)


def bullets(items):
    for item in items:
        par = doc.add_paragraph(style="List Bullet")
        if isinstance(item, str):
            par.add_run(item)
        else:
            for text, kind in item:
                run = par.add_run(text)
                if kind == "b":
                    run.bold = True
                elif kind == "code":
                    run.font.name = "Consolas"
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = ACCENT
        par.paragraph_format.space_after = Pt(3)


def code(text):
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        par = doc.add_paragraph()
        run = par.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = INK
        pf = par.paragraph_format
        pf.space_before = Pt(6 if i == 0 else 0)
        pf.space_after = Pt(6 if i == len(lines) - 1 else 0)
        pf.left_indent = Inches(0.22)
        pf.line_spacing = 1.0
        shade(par, CODE_BG)


def note(label, text):
    par = doc.add_paragraph()
    r1 = par.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = ACCENT
    par.add_run(text)
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(10)
    par.paragraph_format.left_indent = Inches(0.12)
    shade(par, NOTE_BG)
    border(par, "1B4DD1", "left", 18)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(10.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            if val.startswith("`") and val.endswith("`"):
                run = par.add_run(val.strip("`"))
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            else:
                run = par.add_run(val)
                run.font.size = Pt(10.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def rule():
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(4)
    border(par, "D3CBB9", "bottom", 8)


# ======================================================================= COVER
title = doc.add_paragraph("Portfolio Editing Guide", style="Title")
title.paragraph_format.space_after = Pt(2)
sub = p("Ryan Nepal — Mechanical Engineering Portfolio", size=13, color=MUTE)
sub.paragraph_format.space_after = Pt(2)
p("How to add photos, edit text, save your work, and publish it on GitHub.",
  size=11.5, color=MUTE)
rule()

p("You do not need to know how to code to use this. Every word, link and photo "
  "filename on your website comes from a single file, and this guide walks you "
  "through changing it.")

note("Start here",
     "Your entire website is controlled by one file called content.json. You will "
     "never need to open the HTML, CSS or JavaScript files. If you only remember "
     "one thing from this guide, remember that.")

# =============================================================== 1. THE BASICS
h("1. How your website is organised", 1)

p("Your project folder contains these things. Only the first one matters day to day.")

table(
    ["Item", "What it is"],
    [
        ["`content.json`", "All of your content. The only file you edit."],
        ["`images/profile/`", "Your headshot and your About photo."],
        ["`images/projects/`", "Every project photo."],
        ["`files/`", "Your resume PDF."],
        ["`index.html`", "The front page. Leave it alone."],
        ["`project.html`", "The project detail pages. Leave it alone."],
        ["`assets/`", "Styling and code. Only touched to change colours."],
    ],
    widths=[1.7, 4.4],
)

h("What is inside content.json", 2)

p("The file is split into eight labelled blocks, in the same order they appear on "
  "your website. Find the block, change the text, save.")

table(
    ["Block", "Controls"],
    [
        ["`settings`", "Site-wide on/off switches."],
        ["`profile`", "Your name, tagline, hero buttons, the fact strip."],
        ["`about`", "The About paragraphs and the three highlight numbers."],
        ["`projects`", "Every project, its photos and its detail page."],
        ["`experience`", "Jobs, internships and club roles."],
        ["`skills`", "Skill categories and the skills inside them."],
        ["`education`", "Degree, school, graduation date, coursework."],
        ["`contact`", "Email, phone, LinkedIn, GitHub, location."],
    ],
    widths=[1.4, 4.7],
)

h("The four rules of editing content.json", 2)

steps([
    [("Text goes inside double quote marks. ", ""),
     ('"jobTitle": "Design Engineer"', "code")],
    [("A comma goes ", ""), ("between", "b"),
     (" items, but ", ""), ("never", "b"),
     (" before a closing ", ""), ("}", "code"), (" or ", ""), ("]", "code"), (".", "")],
    [("Use straight quote marks only. Microsoft Word silently turns ", ""),
     ('"straight quotes"', "code"), (" into ", ""), ("“curly quotes”", "code"),
     (", which breaks the file. If you draft text in Word, paste it into Notepad "
      "first, then copy it from there.", "")],
    [("Anything starting with an underscore, like ", ""), ('"_note"', "code"),
     (", is a comment written for you. The website ignores it. Leave it or delete "
      "it, either is fine.", "")],
])

note("Deleting is safe",
     "If you do not want a button, a skill, a job or a whole project, delete its "
     "block. The website only draws what exists. Nothing will look broken or leave "
     "an empty gap.")

note("The word TODO is special",
     "Any link whose address starts with TODO is hidden automatically, so an "
     "unfinished link never appears as a broken button. TODO inside ordinary text "
     "does show on the page — that is deliberate, so you can spot what still needs "
     "writing.")

# ================================================================== 2. PHOTOS
h("2. Adding photos", 1)

p("Adding a photo is always the same two steps: put the file in the right folder, "
  "then write its file name into content.json.")

h("Where each photo goes", 2)

table(
    ["Photo", "Folder", "File name to use"],
    [
        ["Your headshot, top of the page", "`images/profile/`", "`headshot.jpg`"],
        ["The photo beside About", "`images/profile/`", "`training-wheels.jpg`"],
        ["Any project photo", "`images/projects/`", "your choice"],
    ],
    widths=[2.3, 1.7, 2.1],
)

h("Your headshot — the easiest one", 2)

steps([
    [("Save your photo with the exact name ", ""), ("headshot.jpg", "code")],
    [("Put it in the ", ""), ("images/profile/", "code"), (" folder", "")],
    "That is all. content.json already points at that name, so there is nothing to edit.",
])

p("The About photo works the same way. It is currently training-wheels.jpg in that "
  "same folder. To swap it for a different photo, either give your new file that "
  "exact name, or put your file in the folder and change the photo line inside the "
  "about block of content.json.")

h("A project cover photo", 2)

p("This is the photo on the project card on the front page.")

steps([
    [("Put your photo in ", ""), ("images/projects/", "code")],
    [("Open ", ""), ("content.json", "code"),
     (" and find the project. Change the coverPhoto line to your file name:", "")],
])

code('"coverPhoto": "mazda3-cover.jpg",')

p("Your three projects currently expect these names, so if you use them there is "
  "nothing to edit at all:")

bullets([
    [("supermileage-fuel-system-cover.jpg", "code")],
    [("leak-test-fixture-cover.jpg", "code")],
    [("mazda3-cover.jpg", "code")],
])

h("Extra photos on a project's detail page", 2)

p("Each project has a photos list that is currently empty. Fill it in and the "
  "gallery builds itself. Captions are optional but worth writing.")

code("""
"photos": [
  { "file": "mazda3-01.jpg", "caption": "Front end reassembled" },
  { "file": "mazda3-02.jpg", "caption": "Tapped manifold ports" },
  { "file": "mazda3-03.jpg", "caption": "Vacuum lines installed" }
],
""")

p("The layout adapts on its own: one photo fills the width, two sit side by side, "
  "three or more form a grid. You do not need to do anything different.")

h("Naming and sizing photos", 2)

p("Two habits that prevent problems later.")

h("Names", 3)
p("Lowercase letters, numbers and hyphens only. No spaces, no brackets, no # symbol.")
bullets([
    [("Good:  ", "b"), ("mazda3-01.jpg", "code"), ("   ", ""),
     ("fuel-pump-adapter-cover.jpg", "code")],
    [("Bad:  ", "b"), ("IMG_2841 (1).JPG", "code"), ("   ", ""), ("My Photo #2.jpg", "code")],
])

h("Sizes", 3)
table(
    ["Photo", "Long edge", "Keep under"],
    [
        ["Headshot", "1000 pixels", "300 KB"],
        ["About photo", "1200 pixels", "300 KB"],
        ["Project cover", "1600 pixels", "400 KB"],
        ["Gallery photo", "1600 pixels", "400 KB"],
    ],
    widths=[2.3, 1.9, 1.9],
)

p("Photos straight off a phone are often 5 to 10 megabytes, which makes your site "
  "slow to load. Shrinking them is quick and free:")

steps([
    [("Go to ", ""), ("squoosh.app", "b")],
    "Drag your photo onto the page",
    "On the right, set the width to 1600 and drag quality down to about 70",
    "Click the download arrow",
])

p("A 6 MB photo becomes roughly 250 KB and looks identical on a screen.")

h("Describing photos for screen readers", 2)

p("Every photo should have a short factual description so the site works for "
  "visitors using a screen reader. One sentence is enough.")

code('"alt": "Aluminium adapter clamped in the mill vise",')

p('Write what is actually in the photo. Not "photo", not "image1".')

note("Missing photos never break anything",
     "Until a photo file exists, your site draws a neutral dashed placeholder box "
     "in its place. That is why the site looks finished right now even though you "
     "have not added a single photo. Publish first, add photos whenever you like.")

# ================================================================ 3. PROJECTS
h("3. Projects", 1)

h("Adding a new project", 2)

p("There is a ready-made template waiting for you at the bottom of the projects "
  "list. It is invisible on your website and exists purely to be copied.")

steps([
    [("Open ", ""), ("content.json", "code"),
     (" and scroll to the bottom of the projects list", "")],
    [("Find the block that begins with ", ""), ('"hidden": true', "code"),
     (" and is labelled TEMPLATE", "")],
    [("Select the whole block, from its opening ", ""), ("{", "code"),
     (" to its closing ", ""), ("}", "code"), (", and copy it", "")],
    [("Paste it ", ""), ("above", "b"),
     (" the project you want it to appear after, and put a comma between the two "
      "blocks", "")],
    [("Change ", ""), ('"hidden": true', "code"), (" to ", ""),
     ('"hidden": false', "code")],
    "Fill in your text and delete any lines you do not need",
])

h("What each field does", 2)

table(
    ["Field", "What it does"],
    [
        ["`title`", "The project name."],
        ["`blurb`", "Two or three sentences. This is the text on the card."],
        ["`date`", "Shown on the detail page."],
        ["`setting`", "Coursework, internship, personal project, club."],
        ["`role`", "What you personally did."],
        ["`status`", "Complete, in progress, built and tested."],
        ["`disciplines`", "The small blue words above the title."],
        ["`tools`", "The little tags on the card. First six appear."],
        ["`coverPhoto`", "The photo on the card."],
        ["`stats`", "The big coloured numbers on the detail page."],
        ["`problem`", "What needed solving, and why it was hard."],
        ["`approach`", "One bullet per step you took."],
        ["`results`", "What came out of it. Numbers if you have them."],
        ["`photos`", "The gallery on the detail page."],
        ["`links`", "Buttons for a report, CAD file or repository."],
        ["`id`", "The web address of the page. Lowercase and hyphens only."],
    ],
    widths=[1.5, 4.6],
)

note("Numbering is automatic",
     "The 01, 02, 03 labels follow the order of the list. To reorder your "
     "projects, move whole blocks up or down. The numbers renumber themselves.")

h("Removing a project", 2)

p("Two ways, depending on whether you might want it back.")

bullets([
    [("Hide it: ", "b"), ("add ", ""), ('"hidden": true,', "code"),
     (" as the first line inside the project. It stays in the file but disappears "
      "from the website.", "")],
    [("Delete it: ", "b"), ("remove the whole block, from its opening ", ""),
     ("{", "code"), (" to its closing ", ""), ("}", "code"),
     (", plus the comma that follows it.", "")],
])

h("The filter buttons", 2)

p("There is a row of filter buttons that can sit above your projects. It is "
  "switched off, because filtering three projects is not worth the space. Once you "
  "have six or so, turn it on at the very top of content.json:")

code("""
"settings": {
  "showProjectFilters": true
}
""")

p("No quote marks around true or false. The buttons build themselves from the "
  "disciplines you listed on each project, so there is nothing else to maintain.")

# ============================================================== 4. EXPERIENCE
h("4. Jobs, internships and club roles", 1)

p("These live in the experience block. Newest goes at the top.")

steps([
    "Copy an existing job block",
    "Paste it where you want it in the list",
    "Put a comma between the two blocks",
    "Change the text",
])

code("""
{
  "jobTitle": "Mechanical Design Engineer Co-op",
  "organization": "Some Company, Inc.",
  "dates": "January 2027 — June 2027",
  "location": "Boston, MA",
  "kind": "Co-op",
  "bullets": [
    "What you did, ending with the result.",
    "Second accomplishment.",
    "Third accomplishment."
  ]
},
""")

bullets([
    [("kind", "code"),
     (" is the small tag beside the dates. Write anything: Internship, Co-op, "
      "Leadership, Research, Part-time. Leave the line out and no tag appears.", "")],
    "Two to four bullets each reads best. Lead with what you did, end with the result.",
])

# ================================================================== 5. SKILLS
h("5. Skills", 1)

p("Skills are grouped into categories. Each category looks like this:")

code("""
{
  "group": "CAD & Design",
  "items": [
    { "name": "SolidWorks", "note": "parts, assemblies, drawings" },
    "Autodesk Inventor"
  ]
}
""")

p("There are two ways to write a skill, and you can mix them freely in the same list:")

bullets([
    [("Just the name: ", ""), ('"Mastercam"', "code")],
    [("Name plus a small grey note: ", ""),
     ('{ "name": "Mastercam", "note": "2D contour, pocketing" }', "code")],
])

h("To add a skill", 2)
p("Add a line to the items list, with a comma after the line above it.")

h("To add a whole new category", 2)
p("Copy a full group block, paste it, and change the group name and items. "
  "Categories appear in the order they are listed.")

h("To remove a skill or category", 2)
p("Delete its line, or its whole block.")

# =============================================================== 6. EDUCATION
h("6. Education", 1)

p("A short block. Change the text between the quote marks.")

code("""
"education": {
  "degree": "Bachelor of Science in Mechanical Engineering",
  "school": "University of Massachusetts Amherst",
  "location": "Amherst, MA",
  "graduation": "Expected May 2027",
  "coursework": [
    "Design of Mechanical Components",
    "Fluid Mechanics"
  ]
}
""")

p("To add a course, add a line to the coursework list with a comma after the line "
  "above it.")

h("Adding honours or certifications later", 2)

p("These sections are not on your site at the moment. They only appear when they "
  "exist, so to bring one back, paste it into the education block:")

code("""
"honors": [
  "Dean's List, Fall 2026"
],
"certifications": [
  { "name": "SolidWorks CSWA", "issuer": "Dassault Systèmes", "year": "2026" }
],
""")

# ================================================================= 7. CONTACT
h("7. Contact details", 1)

code("""
"contact": {
  "heading": "Open to mechanical engineering internships and co-ops.",
  "email": "nepalryan@outlook.com",
  "phone": "617-396-1412",
  "linkedin": "https://www.linkedin.com/in/ryan-nepal",
  "linkedinLabel": "linkedin.com/in/ryan-nepal",
  "github": "TODO: paste your GitHub profile URL here",
  "githubLabel": "",
  "resume": "files/Ryan-Nepal-Resume.pdf",
  "location": "Cambridge, MA"
}
""")

bullets([
    [("heading", "code"),
     (" is the sentence under GET IN TOUCH. Update it when what you are looking "
      "for changes.", "")],
    [("The ", ""), ("Label", "code"),
     (" versions are what visitors see. The plain versions are where the link "
      "actually goes. So linkedin is the full https address, and linkedinLabel is "
      "the short tidy version printed on the page.", "")],
    [("To add your GitHub: ", "b"),
     ("paste your profile address over the TODO in github, then put "
      "github.com/yourusername in githubLabel. A GitHub button and a GitHub row "
      "both appear on their own.", "")],
    [("To remove a row: ", "b"),
     ("delete its line. Deleting the phone line removes your phone number from the "
      "page entirely, which is worth considering on a public website.", "")],
])

p("There is no message form, by design. The Email me button opens the visitor's own "
  "mail app, so there is no third-party service to sign up for, pay for, or have "
  "break quietly while you are not looking.")

# ================================================================== 8. RESUME
h("8. Replacing your resume", 1)

h("The easy way", 2)
steps([
    [("Name your new PDF exactly ", ""), ("Ryan-Nepal-Resume.pdf", "code")],
    [("Put it in the ", ""), ("files/", "code"),
     (" folder, replacing the one already there", "")],
    "Done. Nothing in content.json needs to change.",
])

h("If you want a different file name", 2)
p("Put your PDF in the files folder, then change it in two places in content.json:")

code("""
"buttons": [
  { "label": "Download resume", "url": "files/My-New-Resume.pdf", "primary": true }
]

"contact": {
  "resume": "files/My-New-Resume.pdf"
}
""")

# =================================================================== 9. COLOUR
h("9. Changing the colour", 1)

p("This is the one time you open a file other than content.json, and it is a "
  "single line.")

steps([
    [("Open ", ""), ("assets/css/styles.css", "code")],
    "Look near the very top, around line 22",
    [("Change the six characters after the ", ""), ("#", "code")],
])

code("--accent: #1B4DD1;")

p("That one line recolours your name, every link, every number, the active menu "
  "item and every button. Options that keep text readable:")

table(
    ["Colour", "Code"],
    [
        ["Blue (current)", "`#1B4DD1`"],
        ["Deep green", "`#1F6F4A`"],
        ["Rust orange", "`#A6410E`"],
        ["Indigo", "`#4B3BC4`"],
        ["Teal", "`#0F6E7A`"],
        ["Oxblood", "`#8E1F3C`"],
    ],
    widths=[2.4, 1.8],
)

p("Two more lines nearby you might want. The first is the page background, the "
  "second is how much vertical space sits between sections.")

code("""
--paper: #F4F1EA;
--section-space: 128px;
""")

note("Avoid pale colours",
     "Yellows and light greens fail readability standards for small text. Stick to "
     "something with real depth to it.")

# ================================================================ 10. PREVIEW
h("10. Previewing on your own computer", 1)

p("Optional, but useful before making a big change. Note that double-clicking "
  "index.html will not work: web browsers refuse to read content.json from a file "
  "opened directly off your hard drive. You need a small local server, which is "
  "one line.")

steps([
    [("Open a terminal in your project folder, ", ""),
     ("C:\\Users\\rnep\\EngPro", "code")],
    "Type this line and press Enter:",
])

code("py -m http.server 5180 --bind 127.0.0.1")

steps([
    [("Open ", ""), ("http://127.0.0.1:5180", "b"), (" in your browser", "")],
    [("Leave that terminal window open while you browse. Press ", ""),
     ("Ctrl + C", "b"), (" in it when you are finished", "")],
])

bullets([
    [("On Windows the command is ", ""), ("py", "code"), (", not ", ""),
     ("python", "code"), (".", "")],
    [("Keep the ", ""), ("--bind 127.0.0.1", "code"),
     (" part. Without it you may end up looking at a different website by "
      "accident if something else on your computer is using that port.", "")],
    "If you edit a file while the browser is open, press Ctrl + Shift + R to force "
    "it to reload properly.",
])

note("You may never need this",
     "Once your site is on GitHub, the easiest way to work is to edit content.json "
     "in your web browser on github.com and look at the live site. No terminal, no "
     "server, nothing installed.")

# ================================================================= 11. DEPLOY
h("11. Publishing on GitHub — first time only", 1)

p("This takes about ten minutes and you only do it once. GitHub Pages hosting is "
  "free and gives you a real web address for your resume.")

h("Step 1 — Make a GitHub account", 2)
p("Go to github.com/signup. Your username becomes part of your web address, so "
  "pick something you would be happy putting on a resume.")

h("Step 2 — Create the repository", 2)
steps([
    "Go to github.com/new",
    [("Repository name: ", ""), ("portfolio", "code")],
    [("Choose ", ""), ("Public", "b"),
     (". Free hosting only works on public repositories.", "")],
    "Leave the boxes for README, .gitignore and license unticked",
    [("Click ", ""), ("Create repository", "b")],
])

h("Step 3 — Upload your files", 2)
steps([
    [("On the empty repository page, click ", ""),
     ("uploading an existing file", "b")],
    [("Open ", ""), ("C:\\Users\\rnep\\EngPro", "code"),
     (" on your computer", "")],
    "Select everything and drag it all into the browser window",
    [("In the box at the bottom type ", ""), ("initial upload", "code"),
     (", then click ", ""), ("Commit changes", "b")],
])

p("Two things to watch:")

bullets([
    [("Drag the folders themselves", "b"),
     (", not just the files inside them. The assets, images and files folders must "
      "keep their structure.", "")],
    [(".nojekyll", "code"),
     (" is a real file with no name before the dot, and it must be uploaded. If "
      "Windows hides it, go to File Explorer, then View, then tick Hidden items.", "")],
])

p("You can delete the old Post Summer - Ryan Nepal (1).pdf from the folder before "
  "uploading. The copy your website actually uses is files/Ryan-Nepal-Resume.pdf.")

h("Step 4 — Switch on GitHub Pages", 2)
steps([
    [("In your repository, click ", ""), ("Settings", "b"),
     (" in the top row", "")],
    [("In the left sidebar, click ", ""), ("Pages", "b")],
    [("Set Source to ", ""), ("Deploy from a branch", "b")],
    [("Set Branch to ", ""), ("main", "code"), (" and Folder to ", ""),
     ("/ (root)", "code")],
    [("Click ", ""), ("Save", "b")],
])

h("Step 5 — Find your web address", 2)
p("Wait a minute or two and refresh that same Settings page. A green banner appears "
  "with your address:")
code("https://yourusername.github.io/portfolio/")
p("That is your portfolio. Put it on your resume and your LinkedIn profile.")

note("A 404 at first is normal",
     "The first publish takes one to three minutes. Wait, then reload. You can "
     "watch progress under the Actions tab — a green checkmark means it is live.")

h("Step 6 — Check it over", 2)
bullets([
    "Open it on your laptop and on your phone",
    "Does your name and the resume button appear?",
    "Do the three project cards appear, and does clicking one open its detail page?",
    "Does every menu item jump to the right section?",
    "Does the resume button download your PDF?",
])
p("Photos will still show dashed placeholder boxes until you upload them. That is "
  "expected.")

# ============================================================== 12. UPDATING
h("12. Saving and republishing after an edit", 1)

note("There is no separate publish step",
     "Saving is publishing. On GitHub, saving is called Commit changes, and "
     "committing automatically rebuilds your live website. About a minute later it "
     "is public. That is the whole system.")

h("The everyday routine", 2)

p("This works from a laptop or a phone, in a web browser, with nothing installed.")

steps([
    [("Go to ", ""), ("github.com/yourusername/portfolio", "code")],
    [("Click the file you want to change, almost always ", ""),
     ("content.json", "code")],
    [("Click the ", ""), ("pencil icon", "b"),
     (" at the top right of the file. On a phone, tap the three dots, then Edit.", "")],
    "Make your edit",
    [("Scroll to the bottom and type a short note describing the change, such as ", ""),
     ("add spring co-op", "code")],
    [("Click ", ""), ("Commit changes", "b")],
    "Wait about a minute, then reload your live site",
])

h("Uploading a photo", 2)
steps([
    [("Click into the ", ""), ("images/projects/", "code"), (" folder", "")],
    [("Click ", ""), ("Add file", "b"), (", then ", ""), ("Upload files", "b")],
    "Drag your photo in, or tap choose your files on a phone",
    [("Click ", ""), ("Commit changes", "b")],
    [("Then edit ", ""), ("content.json", "code"),
     (" to write in the new file name", "")],
])

h("Updating your resume", 2)
steps([
    [("Click into the ", ""), ("files/", "code"), (" folder", "")],
    [("Click ", ""), ("Add file", "b"), (", then ", ""), ("Upload files", "b")],
    [("Upload your new PDF named exactly ", ""),
     ("Ryan-Nepal-Resume.pdf", "code")],
    [("Click ", ""), ("Commit changes", "b"),
     (". It replaces the old one and nothing else needs changing.", "")],
])

note("One habit worth forming",
     "After you publish, treat GitHub as the real copy of your website and make "
     "your edits there. The folder on your computer and the one on GitHub are "
     "separate. If you edit both, they drift apart and you will have to work out "
     "which one is current. Editing in the browser avoids this completely.")

# =============================================================== 13. TROUBLE
h("13. If something goes wrong", 1)

p("Your website cannot be permanently broken. Every previous version is saved and "
  "you can always go back.")

h('The page says "content.json has a typo in it"', 2)

p("That message is your site telling you exactly what is wrong. In order of how "
  "often it happens:")

steps([
    "A missing comma between two items",
    "An extra comma just before a closing } or ]",
    "Curly quotes pasted in from Word",
    "A missing closing brace or bracket after copying a block",
])

p("The extra comma is the easiest to miss. This is broken:")

code("""
"bullets": [
  "First bullet",
  "Second bullet",
]
""")

p("The comma after the last item has nothing following it, so remove it.")

p("To find any of these in about ten seconds: go to jsonlint.com, paste in the "
  "whole contents of content.json, and click Validate JSON. It names the exact line.")

h("A photo shows a dashed Image not found box", 2)
p("The file name in content.json does not match the file in your images folder. "
  "Check for:")
bullets([
    "Capital letters. On a web server, .JPG and .jpg are different files.",
    "A space in the file name.",
    "The photo being in images/profile/ when the site is looking in images/projects/.",
])

h("The page is completely blank", 2)
p("A file was probably renamed or moved. Check that content.json, index.html and "
  "the assets folder are all at the top level of your repository, and that "
  "content.json is spelled exactly that way, all lowercase.")

h("The site loads but has no styling", 2)
p("The stylesheet is not where the page expects it. In your repository you should "
  "be able to click assets, then css, then styles.css. If those folders are "
  "missing, upload them again.")

h("I edited a file but the site looks the same", 2)
bullets([
    "It may still be building. Check the Actions tab for a running job.",
    "Or your browser cached the old version. Press Ctrl + Shift + R on Windows, or "
    "Cmd + Shift + R on a Mac. On a phone, open the site in a private tab.",
])

h("I broke something and want the old version back", 2)
steps([
    [("In your repository, click ", ""), ("Commits", "b"),
     (" — the clock icon above the file list", "")],
    "Find the commit from before the one that broke things",
    [("Click the ", ""), ("...", "code"), (" next to the bad commit, then ", ""),
     ("Revert", "b")],
])
p("Or simply edit the file again and fix the mistake. Either works. Nothing is ever "
  "lost.")

# ================================================================ QUICK TABLE
doc.add_page_break()
h("Quick reference", 1)

table(
    ["I want to change...", "Go to"],
    [
        ["My name, tagline, hero buttons, the fact strip", "content.json → profile"],
        ["The About paragraphs and highlight numbers", "content.json → about"],
        ["Anything about a project", "content.json → projects"],
        ["Jobs, internships, club roles", "content.json → experience"],
        ["Skill categories and skills", "content.json → skills"],
        ["Degree, school, graduation, coursework", "content.json → education"],
        ["Email, phone, LinkedIn, GitHub, location", "content.json → contact"],
        ["Show or hide the project filter buttons", "content.json → settings"],
        ["The accent colour, background or spacing", "assets/css/styles.css, top of file"],
        ["My headshot", "replace images/profile/headshot.jpg"],
        ["My resume", "replace files/Ryan-Nepal-Resume.pdf"],
    ],
    widths=[3.4, 2.8],
)

h("Things still marked TODO in your content", 1)

p("Search content.json for the word TODO and you will find these waiting for you:")

bullets([
    "Your GitHub profile address, in two places",
    "One About paragraph, to write in your own voice",
    "Programming languages under Skills",
    "Photos for the three reserved slots on the leak test fixture project",
    "The engineering drawing and CAD render for the two reserved slots on the "
    "fuel system project",
    "Photo descriptions, once you add the photos",
    "An optional CAD file link on the Supermileage project",
])

rule()
final = p("Ryan Nepal — Mechanical Engineering Portfolio · Editing Guide",
          size=9.5, color=MUTE)
final.alignment = WD_ALIGN_PARAGRAPH.CENTER

import sys

out = sys.argv[1] if len(sys.argv) > 1 else "Portfolio Editing Guide.docx"
doc.save(out)
print("Saved:", out)
